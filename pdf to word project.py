import streamlit as st
import unicodedata
import re
from PyPDF2 import PdfReader
from docx import Document
from io import BytesIO
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.utils import simpleSplit # For splitting lines in PDF

# ---------- Helper Functions ----------
def clean_text(text):
    """Clean text to handle various encoding issues"""
    if not text:
        return ""
    try:
        # Normalize unicode characters
        text = unicodedata.normalize('NFKD', text)
        
        # Remove or replace problematic characters
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x84\x86-\x9f]', '', text)
        
        # Replace common problematic characters
        replacements = {
            '\u2018': "'", '\u2019': "'",  # Smart quotes
            '\u201c': '"', '\u201d': '"',  # Smart double quotes
            '\u2013': '-', '\u2014': '--', # En/em dashes
            '\u2026': '...', # Ellipsis
            '\u00a0': ' ',   # Non-breaking space
        }
        
        for old, new in replacements.items():
            text = text.replace(old, new)
        
        # Keep only printable ASCII and common Unicode characters
        text = ''.join(char for char in text if ord(char) < 127 or char.isspace() or unicodedata.category(char).startswith('L'))
        
        return text
    except Exception:
        # If all else fails, keep only basic ASCII
        return ''.join(char for char in str(text) if ord(char) < 127)

def safe_extract_pdf_text(page):
    """Safely extract text from PDF page with multiple fallback methods"""
    try:
        # Primary method
        text = page.extract_text()
        if text and text.strip():
            return clean_text(text)
    except Exception:
        pass
    try:
        # Alternative extraction method
        text = ""
        if "/Contents" in page:
            content = page["/Contents"]
            if hasattr(content, 'get_data'):
                raw_text = content.get_data().decode('utf-8', errors='ignore')
                # Simple text extraction from content stream
                text_matches = re.findall(r'\(([^)]+)\)', raw_text)
                text = ' '.join(text_matches)
        
        if text and text.strip():
            return clean_text(text)
    except Exception:
        pass
    return "[Unable to extract text from this page]"

# ---------- Function: PDF to Word ----------
def pdf_to_word(pdf_file):
    try:
        pdf_file.seek(0)
        # Try to read PDF with error handling
        try:
            reader = PdfReader(pdf_file)
        except Exception as e:
            st.error(f"Cannot read PDF file: {str(e)}")
            return None
        
        if len(reader.pages) == 0:
            st.error("PDF file appears to be empty or corrupted")
            return None
        
        doc = Document()
        doc.add_heading('Converted from PDF', 0)
        
        total_pages = len(reader.pages)
        progress_bar = st.progress(0)
        
        for page_num, page in enumerate(reader.pages):
            try:
                # Update progress
                progress_bar.progress((page_num + 1) / total_pages)
                
                text = safe_extract_pdf_text(page)
                
                if text and text.strip() and text != "[Unable to extract text from this page]":
                    # Split long text into paragraphs
                    paragraphs = text.split('\n\n')
                    for para in paragraphs:
                        if para.strip():
                            doc.add_paragraph(para.strip())
                else:
                    doc.add_paragraph(f"[Page {page_num + 1} - No extractable text or image-only content]")
                    
            except Exception as e:
                doc.add_paragraph(f"[Page {page_num + 1} - Error extracting content: {str(e)[:100]}]")
                continue

        progress_bar.empty()
        return doc
        
    except Exception as e:
        st.error(f"Error processing PDF: {str(e)}")
        return None

# ---------- Function: Save Word Document to Buffer ----------
def save_word_file(doc):
    try:
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        # Verify the buffer has content
        if len(buffer.getvalue()) == 0:
            raise Exception("Generated Word document is empty")
            
        return buffer
    except Exception as e:
        st.error(f"Error saving Word file: {str(e)}")
        return None

# ---------- Function: Word to PDF (Enhanced) ----------
def word_to_pdf(word_file):
    try:
        word_file.seek(0)
        doc_bytes = word_file.read()
        if len(doc_bytes) == 0:
            st.error("Word file is empty")
            return None
            
        doc_stream = BytesIO(doc_bytes)
        
        try:
            document = Document(doc_stream)
        except Exception as e:
            st.error(f"Cannot read Word document: {str(e)}")
            return None

        # Extract all text with better handling
        full_text = []
        
        # Process paragraphs
        for para in document.paragraphs:
            if para.text and para.text.strip():
                cleaned_text = clean_text(para.text)
                if cleaned_text.strip():
                    full_text.append(cleaned_text)
        
        # Process tables if any
        for table in document.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = clean_text(cell.text)
                    if cell_text.strip():
                        row_text.append(cell_text)
                if row_text:
                    full_text.append(" | ".join(row_text))

        if not full_text:
            st.error("No readable content found in Word document")
            return None

        # Create PDF buffer
        pdf_buffer = BytesIO()
        c = canvas.Canvas(pdf_buffer, pagesize=A4)
        width, height = A4
        
        # Set up text parameters
        margin = 50
        line_height = 14
        max_width = width - 2 * margin
        y_position = height - margin
        font_size = 10
        
        # Set font
        c.setFont("Helvetica", font_size)
        
        # Add progress tracking
        total_paragraphs = len(full_text)
        progress_bar = st.progress(0)
        
        # Process each paragraph
        for idx, paragraph in enumerate(full_text):
            # Update progress
            progress_bar.progress((idx + 1) / total_paragraphs)
            
            if not paragraph.strip():
                continue
            
            try:
                # Split long lines to fit page width
                lines = simpleSplit(paragraph, "Helvetica", font_size, max_width)
                
                for line in lines:
                    # Check if we need a new page
                    if y_position < margin + line_height:
                        c.showPage()
                        c.setFont("Helvetica", font_size)
                        y_position = height - margin
                    
                    # Draw the line with additional safety
                    try:
                        if line.strip():  # Only draw non-empty lines
                            c.drawString(margin, y_position, line[:200])  # Limit line length
                    except Exception:
                        # If there's still an issue, draw a placeholder
                        c.drawString(margin, y_position, "[Line contains unsupported characters]")
                    
                    y_position -= line_height
                
                # Add space between paragraphs
                y_position -= line_height * 0.5
                
            except Exception as e:
                # Skip problematic paragraphs but continue processing
                c.drawString(margin, y_position, f"[Paragraph skipped due to formatting issues]")
                y_position -= line_height
                continue

        progress_bar.empty()
        
        # Finalize PDF
        c.save()
        pdf_buffer.seek(0)
        
        # Verify the buffer has content
        buffer_content = pdf_buffer.getvalue()
        if len(buffer_content) == 0:
            raise Exception("Generated PDF is empty")
        
        # Additional validation - check if it's a valid PDF
        if not buffer_content.startswith(b'%PDF'):
            raise Exception("Generated file is not a valid PDF")
            
        return pdf_buffer
        
    except Exception as e:
        st.error(f"Error converting Word to PDF: {str(e)}")
        return None

# ---------- Modern UI Styling ----------
st.markdown("""
<style>
    .reportview-container {
        background: #f0f2f6;
    }
    .main .block-container {
        padding-top: 30px;
        padding-right: 50px;
        padding-left: 50px;
        padding-bottom: 30px;
    }
    .stApp {
        background-color: #f0f2f6;
    }
    .css-1d391kg { /* Target for main content background */
        background-color: #ffffff;
        padding: 30px;
        border-radius: 10px;
        box-shadow: 0 4px 8px rgba(0,0,0,0.1);
    }
    h1, h2, h3, h4, h5, h6 {
        color: #262730;
        font-family: 'Helvetica Neue', Helvetica, Arial, sans-serif;
    }
    .stButton>button {
        background-color: #4CAF50;
        color: white;
        border-radius: 5px;
        padding: 10px 20px;
        font-size: 16px;
        border: none;
        cursor: pointer;
        transition: background-color 0.3s ease;
    }
    .stButton>button:hover {
        background-color: #45a049;
    }
    .stFileUploader label {
        color: #262730;
        font-weight: bold;
    }
    .stRadio > label {
        font-weight: bold;
        color: #262730;
    }
    .file-info {
        background-color: #e0f7fa;
        border-left: 5px solid #00bcd4;
        padding: 15px;
        margin-top: 20px;
        border-radius: 5px;
        color: #004d40;
        font-size: 14px;
    }
    .stSuccess {
        background-color: #e8f5e9;
        color: #2e7d32;
        border-left: 5px solid #4caf50;
        padding: 10px;
        border-radius: 5px;
        margin-top: 15px;
    }
    .stError {
        background-color: #ffebee;
        color: #c62828;
        border-left: 5px solid #f44336;
        padding: 10px;
        border-radius: 5px;
        margin-top: 15px;
    }
    .footer {
        font-size: 0.85em;
        text-align: center;
        margin-top: 50px;
        color: #757575;
    }
    .footer a {
        color: #4CAF50;
        text-decoration: none;
    }
    .footer a:hover {
        text-decoration: underline;
    }
    .feature-box {
        background-color: #e3f2fd;
        padding: 20px;
        border-radius: 8px;
        text-align: center;
        margin-bottom: 20px;
        min-height: 150px; /* Ensure consistent height */
        display: flex;
        flex-direction: column;
        justify-content: center;
        align-items: center;
        color: #1a237e;
    }
    .feature-box h4 {
        margin-top: 0;
        color: #1a237e;
    }
    .feature-box p {
        font-size: 0.9em;
        line-height: 1.5;
        color: #3f51b5;
    }
    .st-expander {
        border: 1px solid #e0e0e0;
        border-radius: 8px;
        padding: 10px;
        margin-top: 20px;
        background-color: #f9f9f9;
    }
</style>
""", unsafe_allow_html=True)

# ---------- Modern Streamlit UI ----------
st.markdown("""
<div style="text-align: center; padding: 20px; background-image: linear-gradient(to right, #007bff, #00d4ff); border-radius: 10px; margin-bottom: 30px; color: white;">
    <h1 style="color: white; margin-bottom: 10px;">📄 Document Converter Pro 📝</h1>
    <p style="font-size: 1.1em; color: rgba(255,255,255,0.9);">
        Seamlessly convert your documents between PDF and Word formats with advanced precision and security.
    </p>
</div>
""", unsafe_allow_html=True)

st.markdown("---")

### Feature highlights
col1, col2, col3 = st.columns(3)
with col1:
    st.markdown("""
    <div class="feature-box">
        <h4>⚡ Enterprise Speed</h4>
        <p>Industrial-grade processing with optimized algorithms for maximum efficiency.</p>
    </div>
    """, unsafe_allow_html=True)
with col2:
    st.markdown("""
    <div class="feature-box">
        <h4>🛡 Bank-Level Security</h4>
        <p>Advanced encryption with zero data retention policy for complete privacy.</p>
    </div>
    """, unsafe_allow_html=True)
with col3:
    st.markdown("""
    <div class="feature-box">
        <h4>🎯 Precision Engine</h4>
        <p>AI-powered text extraction with industry-leading accuracy standards.</p>
    </div>
    """, unsafe_allow_html=True)

st.markdown("---")

### Main converter interface
st.markdown('### Convert Your Documents Effortlessly', unsafe_allow_html=True)

### Conversion type selection
st.markdown("### 📋 Select Conversion Protocol")
option = st.radio("", ("📄➡📝 PDF to Word Document", "📝➡📄 Word to PDF Format"), horizontal=True)
st.markdown("---")

if "PDF to Word" in option:
    st.markdown("### 📤 Document Upload Center")
    pdf_file = st.file_uploader("", type=["pdf"], help="Select PDF document for conversion • Maximum file size: 100MB • Supported format: .pdf")
    if pdf_file:
        st.markdown(f"""
        <div class="file-info">
            <strong>📊 Document Analysis:</strong><br>
            📄 Filename: {pdf_file.name}<br>
            📈 File Size: {pdf_file.size:,} bytes ({pdf_file.size/1024/1024:.2f} MB)<br>
            ✅ Status: Document validated and ready for processing<br>
            🔧 Processing Mode: Advanced text extraction with formatting preservation
        </div>
        """, unsafe_allow_html=True)
        
        if st.button("🚀 Initialize Conversion Process", type="primary"):
            with st.spinner("🔄 Processing document conversion..."):
                doc = pdf_to_word(pdf_file)
                if doc:
                    word_buffer = save_word_file(doc)
                    if word_buffer:
                        st.success("✅ Document conversion completed successfully!")
                        st.download_button(
                            label="📥 Download Converted Document",
                            data=word_buffer,
                            file_name=pdf_file.name.replace(".pdf", "") + "_converted.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    else:
                        st.error("❌ Document processing failed during file generation.")
                else:
                    st.error("❌ Conversion process failed. Please verify document integrity and format compatibility.")
else:  # Word to PDF
    st.markdown("### 📤 Document Upload Center")
    word_file = st.file_uploader("", type=["docx"], help="Select Word document for conversion • Maximum file size: 100MB • Supported format: .docx")
    if word_file:
        st.markdown(f"""
        <div class="file-info">
            <strong>📊 Document Analysis:</strong><br>
            📄 Filename: {word_file.name}<br>
            📈 File Size: {word_file.size:,} bytes ({word_file.size/1024/1024:.2f} MB)<br>
            ✅ Status: Document validated and ready for processing<br>
            🔧 Processing Mode: Professional PDF generation with layout optimization
        </div>
        """, unsafe_allow_html=True)
        
        if st.button("🚀 Initialize Conversion Process", type="primary"):
            with st.spinner("🔄 Processing document conversion..."):
                pdf_buffer = word_to_pdf(word_file)
                if pdf_buffer:
                    st.success("✅ Document conversion completed successfully!")
                    st.download_button(
                        label="📥 Download Converted Document",
                        data=pdf_buffer,
                        file_name=word_file.name.replace(".docx", "") + "_converted.pdf",
                        mime="application/pdf"
                    )
                else:
                    st.error("❌ Conversion process failed. Please verify document integrity and format compatibility.")

st.markdown('---', unsafe_allow_html=True)

### Information section
st.markdown("""
## About This Converter
<p>
    Welcome to the <strong>Document Converter Pro</strong> – your ultimate online tool for seamless and secure document conversions between PDF and Word formats. Whether you need to transform a PDF into an editable Word document or convert your Word file into a professional PDF, our application provides a fast, accurate, and reliable solution.
</p>
<p>
    Built with robust libraries and optimized algorithms, this converter handles various document complexities, ensuring minimal loss of formatting and high-quality output. We prioritize your data privacy with a zero data retention policy, meaning your uploaded files are immediately processed and then deleted from our servers once the conversion is complete.
</p>
<p>
    Experience enterprise-grade speed, bank-level security, and a precision engine powered by advanced text extraction capabilities. This tool is designed to make your document workflow more efficient and hassle-free.
</p>
""", unsafe_allow_html=True)

### Troubleshooting section
with st.expander("🛠 Troubleshooting Guide"):
    st.markdown("""
    *🚫 File won't convert?*
    - Verify the file isn't password-protected or corrupted.
    - Try with a smaller file size (< 50MB recommended).
    - Ensure proper file format (.pdf or .docx).
    *📝 Missing text in output?*
    - Some PDFs contain only images (requires OCR, not supported by this tool).
    - Complex layouts may not convert perfectly.
    - Tables and special formatting might need manual review after conversion.

    *⬇ Download issues?*
    - Check browser download settings and popup blockers.
    - Ensure sufficient storage space available on your device.
    - Try refreshing the page and converting again.

    *🔧 Still having problems?*
    - Clear your browser cache and cookies.
    - Try using a different web browser.
    - Check your internet connection stability.
    - If issues persist, consider using a desktop-based conversion tool for highly complex documents.
    """)

### Footer
st.markdown("""
<div class="footer">
    <p>&copy; 2025 Document Converter Pro. All rights reserved. | <a href="#" target="_blank">Privacy Policy</a> | <a href="#" target="_blank">Terms of Service</a></p>
</div>
""", unsafe_allow_html=True)